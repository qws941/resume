#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
import json
import os

def extract_docx_content(file_path):
    """Extract text content from DOCX file"""
    try:
        doc = Document(file_path)
        content = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text.strip())

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_data.append(cell.text.strip())
                if row_data:
                    content.append(" | ".join(row_data))

        return content
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return []

def main():
    resume_files = [
        '/home/jclee/app/resume/이재철_이력서_241217.docx',
        '/home/jclee/app/resume/이재철_이력서_현대오토에버.docx',
        '/home/jclee/app/resume/이재철_이력서_배민.docx',
        '/home/jclee/app/resume/이재철_이력서_당근.docx',
        '/home/jclee/app/resume/이재철_이력서_Toss.docx',
        '/home/jclee/app/resume/이재철_이력서_Coupangpay.docx',
        '/home/jclee/app/resume/이재철_이력서_11번가.docx',
        '/home/jclee/app/resume/이재철 경력사항_프리랜서.docx',
        '/home/jclee/app/resume/이재철_스마일게이트.docx'
    ]

    all_content = {}

    for file_path in resume_files:
        if os.path.exists(file_path):
            filename = os.path.basename(file_path)
            print(f"Processing: {filename}")
            content = extract_docx_content(file_path)
            all_content[filename] = content

    # Save to JSON for analysis
    with open('/home/jclee/app/resume/extracted_content.json', 'w', encoding='utf-8') as f:
        json.dump(all_content, f, ensure_ascii=False, indent=2)

    print("\nExtraction complete! Saved to extracted_content.json")

if __name__ == "__main__":
    main()